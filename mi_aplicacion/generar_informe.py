from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import os
import tempfile
from datetime import datetime
import pythoncom
import win32com.client
import base64

def reemplazar_texto(parrafos, marcador, texto, tamaño_fuente=None):
    for parrafo in parrafos:
        if marcador in parrafo.text:
            nuevo_texto = parrafo.text.replace(marcador, texto)
            parrafo.clear()
            run = parrafo.add_run(nuevo_texto)
            if tamaño_fuente:
                run.font.size = Pt(tamaño_fuente)

def completar_plantilla(usuario, ruta_plantilla, ruta_salida, datos, ruta_qr, foto_base64=None):
    pythoncom.CoInitialize()
    
    try:
        doc = Document(ruta_plantilla)
        fecha_actual = datetime.now().strftime("%Y-%m-%d")
        hora_actual = datetime.now().strftime("%H:%M:%S")

        parrafos = list(doc.paragraphs)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    parrafos.extend(celda.paragraphs)

        for clave, valor in datos.items():
            reemplazar_texto(parrafos, clave, valor)
        reemplazar_texto(parrafos, "{{Fecha}}", fecha_actual, tamaño_fuente=8)
        reemplazar_texto(parrafos, "{{Hora}}", hora_actual, tamaño_fuente=8)
        reemplazar_texto(parrafos,"{{Usuario}}", usuario, tamaño_fuente=8)

        for parrafo in parrafos:
            if "{{QRCode}}" in parrafo.text:
                parrafo.text = parrafo.text.replace("{{QRCode}}", "")
                parrafo.add_run().add_picture(ruta_qr, width=Inches(1.0))
            elif "{{QRCode2}}" in parrafo.text:
                parrafo.text = parrafo.text.replace("{{QRCode2}}", "")
                parrafo.add_run().add_picture(ruta_qr, width=Inches(0.6))

        if foto_base64:
            foto_data = base64.b64decode(foto_base64)
            temp_foto_path = os.path.join(tempfile.gettempdir(), 'temp_foto.png')
            with open(temp_foto_path, 'wb') as temp_foto_file:
                temp_foto_file.write(foto_data)

            for parrafo in parrafos:
                if "{{Foto}}" in parrafo.text:
                    parrafo.text = parrafo.text.replace("{{Foto}}", "")
                    parrafo.add_run().add_picture(temp_foto_path, width=Inches(1.0))

            os.remove(temp_foto_path)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx_path = temp_docx.name
        doc.save(temp_docx_path)

        try:
            convert(temp_docx_path, ruta_salida)
        except Exception as e:
            print(f"Error al convertir DOCX a PDF: {e}")
        finally:
            try:
                word = win32com.client.DispatchEx("Word.Application")
                word.Quit()
            except Exception as e:
                print(f"Error al cerrar Word: {e}")

        os.remove(temp_docx_path)
    
    finally:
        pythoncom.CoUninitialize()
